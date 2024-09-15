import os
import re
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from tkinter import PhotoImage

# Hàm cập nhật nội dung bảng đầu tiên với ngày và serial RTU do người dùng nhập vào mà giữ nguyên định dạng
def update_first_table(doc, test_date, serial_rtu):
    # Lấy bảng đầu tiên từ tài liệu
    first_table = doc.tables[0]

    # Duyệt qua từng hàng trong bảng đầu tiên
    for row in first_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    # Thay thế tag <DATE> bằng giá trị người dùng nhập vào mà giữ nguyên định dạng
                    if "<DATE>" in run.text:
                        run.text = run.text.replace("<DATE>", test_date)
                    # Thay thế tag <SERIAL RTU> bằng giá trị người dùng nhập vào mà giữ nguyên định dạng
                    if "<SERIAL RTU>" in run.text:
                        run.text = run.text.replace("<SERIAL RTU>", serial_rtu)

# Hàm giữ lại các bảng chứa ký tự đặc biệt và xóa các bảng còn lại
def remove_tables_with_special_characters(input_file, output_file, special_tags, test_date, serial_rtu):
    # Mở file Word
    doc = Document(input_file)

    # Cập nhật bảng đầu tiên với thông tin ngày và serial RTU
    update_first_table(doc, test_date, serial_rtu)

    # Duyệt qua từng bảng trong tài liệu và đánh dấu các bảng cần xóa
    tables_to_delete = []
    
    for table_index, table in enumerate(doc.tables):
        # Trường hợp bảng đầu tiên, bảng thứ hai và bảng cuối cùng luôn được giữ lại
        if table_index == 0 or table_index == 1 or table_index == len(doc.tables) - 1:
            continue

        # Kiểm tra xem bảng có đủ 2 hàng hay không (phải có ít nhất 2 hàng để xét hàng thứ hai)
        if len(table.rows) > 1:
            # Lấy nội dung của ô đầu tiên trong hàng thứ hai
            first_cell_text = table.rows[1].cells[0].text.strip()

            # Kiểm tra nếu ô đầu tiên của hàng thứ hai chứa bất kỳ ký tự đặc biệt nào mà người dùng nhập vào
            if any(tag in first_cell_text for tag in special_tags):
                continue  # Giữ lại bảng
            else:
                tables_to_delete.append(table)  # Xóa bảng không có ký tự đặc biệt
        else:
            tables_to_delete.append(table)  # Xóa bảng không đủ hàng

    # Xóa các bảng không nằm trong danh sách giữ lại
    for table in tables_to_delete:
        table._element.getparent().remove(table._element)

    # Lưu file kết quả
    os.makedirs("output", exist_ok=True)  # Tạo thư mục output nếu chưa tồn tại
    doc.save(os.path.join("output", output_file))

# Hàm để cập nhật danh sách file trong thư mục "template"
def update_file_list():
    files = [f for f in os.listdir("template") if f.endswith('.docx')]
    file_list.delete(0, tk.END)  # Xóa danh sách cũ
    for file in files:
        file_list.insert(tk.END, file)

# Hàm xử lý khi người dùng nhấn nút "Xuất văn bản"
def export_file():
    selected_index = file_list.curselection()
    if not selected_index:
        messagebox.showwarning("Chọn file", "Vui lòng chọn một file từ danh sách.")
        return

    input_file = file_list.get(selected_index[0])
    input_path = os.path.join("template", input_file)

    special_tags = special_tag_entry.get().split()
    test_date = date_entry.get()
    serial_rtu = serial_entry.get()

    # Tạo tên file đầu ra
    output_file = f"BienBan_RMU_{input_file.replace('.docx', '')}_{serial_rtu}.docx"

    # Gọi hàm xử lý
    remove_tables_with_special_characters(input_path, output_file, special_tags, test_date, serial_rtu)

    messagebox.showinfo("Thành công", f"File đã được lưu vào thư mục 'output' với tên '{output_file}'")

# Thiết lập GUI
root = tk.Tk()
root.title("Tool tạo biên bản thí nghiệm điện")
root.geometry("400x500")
root.iconbitmap("tool.ico")

# Tác giả phần mềm
author_label = tk.Label(root, text="Phần mềm được tạo bởi Thinhlh", font=("Arial", 10))
author_label.pack(side=tk.BOTTOM, pady=10)

# Label và listbox hiển thị danh sách file trong thư mục "template"
file_label = tk.Label(root, text="Chọn file đầu vào:")
file_label.pack(pady=5)

frame = tk.Frame(root)
frame.pack()

file_list = tk.Listbox(frame, width=50, height=10)
file_list.pack(side=tk.LEFT, pady=5)

# Nút refresh danh sách file
refresh_icon = PhotoImage(file="refresh.png")
refresh_icon = refresh_icon.subsample(20, 20)
refresh_button = tk.Button(frame, image=refresh_icon, command=update_file_list)
refresh_button.pack(side=tk.RIGHT, padx=5)

# Input để nhập ký tự đặc biệt
special_tag_label = tk.Label(root, text="Nhập các ký tự đặc biệt (vd: k1 k2 k3):")
special_tag_label.pack(pady=5)

special_tag_entry = tk.Entry(root, width=50)
special_tag_entry.pack(pady=5)

# Input để nhập ngày thử nghiệm
date_label = tk.Label(root, text="Nhập ngày thử nghiệm (vd: 2024-09-15):")
date_label.pack(pady=5)

date_entry = tk.Entry(root, width=50)
date_entry.pack(pady=5)

# Input để nhập serial RTU
serial_label = tk.Label(root, text="Nhập số serial RTU:")
serial_label.pack(pady=5)

serial_entry = tk.Entry(root, width=50)
serial_entry.pack(pady=5)

# Nút để xuất văn bản
export_button = tk.Button(root, text="Xuất văn bản", command=export_file)
export_button.pack(pady=10)

# Lần đầu tiên cập nhật danh sách file
update_file_list()

# Chạy giao diện GUI
root.mainloop()
