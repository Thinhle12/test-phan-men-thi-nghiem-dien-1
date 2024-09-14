import os
from docx import Document

# Hàm cập nhật nội dung bảng đầu tiên với ngày và serial RTU do người dùng nhập vào mà giữ nguyên định dạng
def update_first_table(doc, test_date, serial_rtu):
    # Lấy bảng đầu tiên từ tài liệu
    first_table = doc.tables[0]

    # Duyệt qua từng hàng trong bảng đầu tiên
    for row in first_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    # Thay thế tag <DATE> bằng giá trị người dùng nhập vào mà giữ nguyên định dạng
                    if "<DATE>" in run.text:
                        run.text = run.text.replace("<DATE>", test_date)
                    # Thay thế tag <SERIAL RTU> bằng giá trị người dùng nhập vào mà giữ nguyên định dạng
                    if "<SERIAL RTU>" in run.text:
                        run.text = run.text.replace("<SERIAL RTU>", serial_rtu)

# Hàm giữ lại các bảng chứa ký tự đặc biệt và xóa các bảng còn lại
def remove_tables_with_special_characters(input_file, output_file, special_tags, test_date, serial_rtu):
    # Mở file Word
    doc = Document(input_file)

    # Cập nhật bảng đầu tiên với thông tin ngày và serial RTU
    update_first_table(doc, test_date, serial_rtu)

    # Duyệt qua từng bảng trong tài liệu và đánh dấu các bảng cần xóa
    tables_to_delete = []
    
    for table_index, table in enumerate(doc.tables):
        print(f"Kiểm tra bảng {table_index + 1}...")

        # Trường hợp bảng đầu tiên, bảng thứ hai và bảng cuối cùng luôn được giữ lại
        if table_index == 0 or table_index == 1 or table_index == len(doc.tables) - 1:
            print(f"Bảng {table_index + 1} là bảng đặc biệt (bảng đầu, bảng thứ hai hoặc bảng cuối cùng) và sẽ được giữ lại.")
            continue

        # Kiểm tra xem bảng có đủ 2 hàng hay không (phải có ít nhất 2 hàng để xét hàng thứ hai)
        if len(table.rows) > 1:
            # Lấy nội dung của ô đầu tiên trong hàng thứ hai
            first_cell_text = table.rows[1].cells[0].text.strip()

            # In ra nội dung của ô đầu tiên trong hàng thứ hai để kiểm tra
            print(f"Nội dung ô đầu tiên của hàng thứ hai: '{first_cell_text}'")

            # Kiểm tra nếu ô đầu tiên của hàng thứ hai chứa bất kỳ ký tự đặc biệt nào mà người dùng nhập vào
            if any(tag in first_cell_text for tag in special_tags):
                print(f"Bảng {table_index + 1} có chứa ký tự đặc biệt và sẽ được giữ lại.")
            else:
                print(f"Bảng {table_index + 1} không chứa ký tự đặc biệt và sẽ bị xóa.")
                tables_to_delete.append(table)
        else:
            print(f"Bảng {table_index + 1} không có đủ 2 hàng, sẽ bị xóa.")
    
    # Xóa các bảng không nằm trong danh sách giữ lại
    for table in tables_to_delete:
        table._element.getparent().remove(table._element)

    # Lưu file kết quả
    doc.save(output_file)
    print(f"Đã lưu kết quả vào '{output_file}'.")

# Hàm liệt kê các file trong thư mục "template" và yêu cầu người dùng chọn file đầu vào
def select_input_file(template_folder):
    # Liệt kê các file trong thư mục "template"
    files = [f for f in os.listdir(template_folder) if f.endswith('.docx')]

    # Hiển thị danh sách file cho người dùng
    print("Chọn file đầu vào từ danh sách:")
    for i, file_name in enumerate(files, start=1):
        print(f"{i}. {file_name}")

    # Yêu cầu người dùng nhập số tương ứng với file
    choice = int(input("Nhập số tương ứng với file bạn muốn chọn: "))
    selected_file = files[choice - 1]

    # Trả về tên file và đường dẫn đầy đủ đến file được chọn
    return selected_file, os.path.join(template_folder, selected_file)

# Hàm tạo thư mục nếu chưa tồn tại
def create_output_folder(output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"Thư mục '{output_folder}' đã được tạo.")
    else:
        print(f"Thư mục '{output_folder}' đã tồn tại.")

# Thư mục chứa các file template
template_folder = 'template'

# Gọi hàm để chọn file đầu vào
selected_file_name, input_file = select_input_file(template_folder)

# Nhập số serial RTU từ người dùng
serial_rtu = input("Nhập số serial RTU: ")

# Đặt tên file đầu ra theo cú pháp yêu cầu
output_file_name = f"BienBan_RMU_{selected_file_name.split('.')[0]}_{serial_rtu}.docx"

# Tạo thư mục "output" nếu chưa tồn tại
output_folder = 'output'
create_output_folder(output_folder)

# Đường dẫn đầy đủ đến file đầu ra trong thư mục "output"
output_file = os.path.join(output_folder, output_file_name)

# Nhập danh sách các ký tự đặc biệt mà người dùng muốn tìm
special_tags = input("Nhập các ký tự đặc biệt, cách nhau bởi khoảng cách (ví dụ: k1 k2 k3): ").split()

# Nhập ngày thử nghiệm từ người dùng
test_date = input("Nhập ngày thử nghiệm (ví dụ: 2024-09-15): ")

# Gọi hàm xử lý
remove_tables_with_special_characters(input_file, output_file, special_tags, test_date, serial_rtu)

print(f"Đã cập nhật bảng đầu tiên và xóa các bảng không có ký tự đặc biệt. Kết quả được lưu vào '{output_file}'.")
